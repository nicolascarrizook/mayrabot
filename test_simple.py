#!/usr/bin/env python3
"""
Simple test script to verify DOCX processing without heavy dependencies
"""

import os
from pathlib import Path

# Check if we can import basic modules
try:
    import docx
    print("✅ python-docx is installed")
except ImportError:
    print("❌ python-docx is NOT installed")
    print("Install with: pip install python-docx")

# Check if DOCX files exist
data_dir = Path("data")
files = ["almuerzoscena.docx", "desayunos.docx", "desayunosequivalentes.docx", "recetas.docx"]

print("\n📁 Checking DOCX files:")
for file in files:
    file_path = data_dir / file
    if file_path.exists():
        size = file_path.stat().st_size / 1024  # KB
        print(f"✅ {file} ({size:.1f} KB)")
    else:
        print(f"❌ {file} NOT FOUND")

# Try to read a sample file
try:
    import docx
    doc = docx.Document(str(data_dir / "almuerzoscena.docx"))
    print(f"\n📄 Sample from almuerzoscena.docx:")
    print(f"   Tables found: {len(doc.tables)}")
    print(f"   Paragraphs found: {len(doc.paragraphs)}")
    
    if doc.tables:
        first_table = doc.tables[0]
        print(f"   First table has {len(first_table.rows)} rows")
except Exception as e:
    print(f"\n❌ Error reading DOCX: {e}")

print("\n✅ Basic test completed!")
print("\nTo run full processing test:")
print("  python data_processor/test_processing.py")
print("\nTo load into ChromaDB (requires .env with OPENAI_API_KEY):")
print("  python data_processor/load_to_chromadb.py")