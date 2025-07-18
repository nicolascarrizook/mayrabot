"""
DOCX Processor Module

This module provides functionality to read and extract data from DOCX files,
specifically designed to handle nutrition recipe documents with tables and text.
"""

import logging
from pathlib import Path
from typing import List, Dict, Any, Optional, Tuple
from dataclasses import dataclass

import docx
from docx.document import Document
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph

logger = logging.getLogger(__name__)


@dataclass
class TableData:
    """Represents extracted table data"""
    headers: List[str]
    rows: List[List[str]]
    metadata: Dict[str, Any]


class DocxProcessor:
    """Base class for processing DOCX files with tables and text"""
    
    def __init__(self, file_path: str):
        """
        Initialize the DOCX processor
        
        Args:
            file_path: Path to the DOCX file
        """
        self.file_path = Path(file_path)
        if not self.file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        self.document: Optional[Document] = None
        self.tables: List[TableData] = []
        self.paragraphs: List[str] = []
        
    def load_document(self) -> None:
        """Load the DOCX document"""
        try:
            self.document = docx.Document(str(self.file_path))
            logger.info(f"Successfully loaded document: {self.file_path}")
        except Exception as e:
            logger.error(f"Error loading document: {e}")
            raise
    
    def extract_tables(self) -> List[TableData]:
        """Extract all tables from the document"""
        if not self.document:
            self.load_document()
        
        self.tables = []
        
        for idx, table in enumerate(self.document.tables):
            table_data = self._process_table(table, idx)
            if table_data:
                self.tables.append(table_data)
        
        logger.info(f"Extracted {len(self.tables)} tables from document")
        return self.tables
    
    def extract_paragraphs(self) -> List[str]:
        """Extract all paragraphs from the document"""
        if not self.document:
            self.load_document()
        
        self.paragraphs = []
        
        for para in self.document.paragraphs:
            text = para.text.strip()
            if text:
                self.paragraphs.append(text)
        
        logger.info(f"Extracted {len(self.paragraphs)} paragraphs from document")
        return self.paragraphs
    
    def _process_table(self, table: Table, index: int) -> Optional[TableData]:
        """
        Process a single table
        
        Args:
            table: The table object to process
            index: Table index in the document
            
        Returns:
            TableData object or None if table is empty
        """
        try:
            # Extract headers from first row
            headers = []
            if len(table.rows) > 0:
                for cell in table.rows[0].cells:
                    headers.append(self._clean_cell_text(cell))
            
            # Extract data rows
            rows = []
            for row_idx, row in enumerate(table.rows[1:], start=1):
                row_data = []
                for cell in row.cells:
                    row_data.append(self._clean_cell_text(cell))
                
                # Skip empty rows
                if any(cell.strip() for cell in row_data):
                    rows.append(row_data)
            
            if headers and rows:
                return TableData(
                    headers=headers,
                    rows=rows,
                    metadata={
                        'table_index': index,
                        'row_count': len(rows),
                        'column_count': len(headers)
                    }
                )
            
            return None
            
        except Exception as e:
            logger.error(f"Error processing table {index}: {e}")
            return None
    
    def _clean_cell_text(self, cell: _Cell) -> str:
        """
        Clean and extract text from a table cell
        
        Args:
            cell: Table cell object
            
        Returns:
            Cleaned text content
        """
        # Extract all paragraphs from the cell
        text_parts = []
        for paragraph in cell.paragraphs:
            text = paragraph.text.strip()
            if text:
                text_parts.append(text)
        
        # Join with space and clean up
        full_text = ' '.join(text_parts)
        # Remove extra whitespace
        full_text = ' '.join(full_text.split())
        
        return full_text
    
    def find_tables_by_header(self, header_pattern: str) -> List[TableData]:
        """
        Find tables that contain specific header pattern
        
        Args:
            header_pattern: Pattern to search in headers
            
        Returns:
            List of matching tables
        """
        if not self.tables:
            self.extract_tables()
        
        matching_tables = []
        for table in self.tables:
            for header in table.headers:
                if header_pattern.lower() in header.lower():
                    matching_tables.append(table)
                    break
        
        return matching_tables
    
    def get_table_as_dict(self, table_data: TableData) -> List[Dict[str, str]]:
        """
        Convert table data to list of dictionaries
        
        Args:
            table_data: TableData object
            
        Returns:
            List of dictionaries with headers as keys
        """
        result = []
        
        for row in table_data.rows:
            row_dict = {}
            for idx, header in enumerate(table_data.headers):
                if idx < len(row):
                    row_dict[header] = row[idx]
                else:
                    row_dict[header] = ""
            
            result.append(row_dict)
        
        return result
    
    def extract_sections(self) -> Dict[str, List[str]]:
        """
        Extract document sections based on heading styles
        
        Returns:
            Dictionary with section titles as keys and content as values
        """
        if not self.document:
            self.load_document()
        
        sections = {}
        current_section = "Introduction"
        current_content = []
        
        for para in self.document.paragraphs:
            # Check if paragraph is a heading
            if para.style.name.startswith('Heading'):
                # Save previous section
                if current_content:
                    sections[current_section] = current_content
                
                # Start new section
                current_section = para.text.strip()
                current_content = []
            else:
                # Add to current section
                text = para.text.strip()
                if text:
                    current_content.append(text)
        
        # Save last section
        if current_content:
            sections[current_section] = current_content
        
        return sections
    
    def search_content(self, keyword: str) -> List[Tuple[str, str]]:
        """
        Search for keyword in document content
        
        Args:
            keyword: Keyword to search for
            
        Returns:
            List of tuples (location, content) containing the keyword
        """
        results = []
        keyword_lower = keyword.lower()
        
        # Search in paragraphs
        if not self.paragraphs:
            self.extract_paragraphs()
        
        for idx, para in enumerate(self.paragraphs):
            if keyword_lower in para.lower():
                results.append((f"Paragraph {idx}", para))
        
        # Search in tables
        if not self.tables:
            self.extract_tables()
        
        for table in self.tables:
            # Search in headers
            for header in table.headers:
                if keyword_lower in header.lower():
                    results.append((
                        f"Table {table.metadata['table_index']} Header",
                        header
                    ))
            
            # Search in rows
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row):
                    if keyword_lower in cell.lower():
                        results.append((
                            f"Table {table.metadata['table_index']} Row {row_idx} Cell {cell_idx}",
                            cell
                        ))
        
        return results